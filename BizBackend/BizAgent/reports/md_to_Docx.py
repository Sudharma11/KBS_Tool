import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def highlight_keywords_in_paragraph(paragraph, text):
    """Adds text to a paragraph, applying bold formatting to keywords and markdown bold syntax."""
    keywords = [
        "Company Overview:", "Founded:", "Headquarters:", "Industry:", "Core Business:",
        "Employee Count:", "Key Executives:", "Competitive Analysis:", "Market Position:",
        "Strengths:", "Weaknesses:", "Opportunities:", "Threats:", "Financial Summary:",
        "Market Size:", "Growth Rate (CAGR):", "Key Trends:", "Recent News Summary:",
        "Source:", "Impact/Sentiment:", "Technology-based SWOT analysis:",
        "Technological strengths:", "Technological weaknesses:"
    ]
    keyword_pattern = '|'.join(re.escape(k) for k in keywords)
    # This pattern captures either **bolded text** or one of the keywords
    pattern = re.compile(f'(\\*\\*.*?\\*\\*|{keyword_pattern})')
    
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        
        is_markdown_bold = part.startswith('**') and part.endswith('**')
        is_keyword = part in keywords

        if is_markdown_bold:
            paragraph.add_run(part.strip('*')).bold = True
        elif is_keyword:
            paragraph.add_run(part).bold = True
        else:
            paragraph.add_run(part)

def convert_md_to_docx(md_path, docx_path, company_name="Target Company", report_type='target'):
    """
    Converts a markdown file to a professionally structured DOCX file
    with a dynamic title page and highlighted keywords.
    """
    if not os.path.exists(md_path):
        print(f"Error: Markdown file not found at {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # --- Create a dynamic title page ---
    if report_type == 'target':
        title_text = f"Target Company Report: {company_name}"
    elif report_type == 'financial':
        title_text = f"Financial Report: {company_name}"
    elif report_type == 'comparison':
        title_text = f"Comparison Report: Kanini vs {company_name}"
    else:
        title_text = "Business Intelligence Report"

    title_p = document.add_paragraph()
    title_p.add_run(title_text).bold = True
    title_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    run = title_p.runs[0]
    run.font.size = Pt(28)

    date_p = document.add_paragraph()
    date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    date_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(50)
    document.add_page_break()

    # --- Process and add the main report content ---
    lines = md_content.split('\n')
    
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue

        # Corrected: More robust filtering for any line that is a TOC heading
        if re.search(r'table of contents', stripped_line, re.IGNORECASE):
            continue

        if stripped_line.startswith('#'):
            level = len(stripped_line) - len(stripped_line.lstrip('#'))
            heading_text = stripped_line.lstrip('# ').strip()
            document.add_heading(heading_text, level=level)
        
        elif stripped_line.startswith('- '):
            item_text = stripped_line.lstrip('- ').strip()
            p = document.add_paragraph(style='List Bullet')
            highlight_keywords_in_paragraph(p, item_text)
        
        else:
            p = document.add_paragraph()
            highlight_keywords_in_paragraph(p, stripped_line)
            p.paragraph_format.space_after = Pt(8)

    try:
        print(f"--- Saving professional DOCX report to: {docx_path} ---")
        document.save(docx_path)
    except Exception as e:
        print(f"Error saving DOCX file at {docx_path}: {e}")



#++++++++++++++++++++++++++++++++++++
# from spire.doc import *
# from spire.doc.common import *

# def convert_md_to_docx(md_file, docx_output):
#     """
#     Convert a Markdown file to DOCX format.
    
#     Args:
#         md_file (str): The path to the Markdown file.
#         docx_output (str): The path for the DOCX output file.
#     """
#     # Create an object of the Document class
#     document = Document()
#     # Load a Markdown file
#     document.LoadFromFile(md_file)

#     # Save the Markdown file to a Word DOCX file
#     document.SaveToFile(docx_output, FileFormat.Docx)
#     # Save the Markdown file to a Word DOC file
#     # document.SaveToFile("MdToDoc.doc", FileFormat.Doc)
#     document.Close()
    
#     print(f"Conversion complete: '{md_file}' converted to '{docx_output}'.")



